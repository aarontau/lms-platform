"""
Convert LMS Project markdown documents to Word (.docx) format.
Uses python-docx with proper heading styles, tables, and code blocks.
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = r"C:\Users\aaron.tau\Desktop\LMS Project"

# ── Helpers ────────────────────────────────────────────────────────────────────

def set_heading_style(paragraph, level):
    """Apply heading style and colour."""
    colours = {
        1: RGBColor(0x1d, 0x4e, 0xd8),   # primary-700 blue
        2: RGBColor(0x15, 0x80, 0x3d),   # secondary-700 green
        3: RGBColor(0x1e, 0x40, 0xaf),   # primary-800
        4: RGBColor(0x37, 0x41, 0x51),   # gray-700
    }
    for run in paragraph.runs:
        run.font.color.rgb = colours.get(level, RGBColor(0, 0, 0))

def add_table_from_md(doc, lines, start_idx):
    """Parse a markdown table and add it to the document. Returns next line index."""
    i = start_idx
    rows = []
    while i < len(lines) and lines[i].strip().startswith('|'):
        row = [cell.strip() for cell in lines[i].strip().strip('|').split('|')]
        if not all(set(c) <= set('-: ') for c in row):   # skip separator rows
            rows.append(row)
        i += 1
    if not rows:
        return i
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = 'Table Grid'
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < col_count:
                cell = row.cells[c_idx]
                cell.text = cell_text
                if r_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    # shade header
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), '1D4ED8')
                    tcPr.append(shd)
    doc.add_paragraph()
    return i

def add_code_block(doc, code_lines):
    """Add a shaded code block."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    run = para.add_run('\n'.join(code_lines))
    run.font.name = 'Courier New'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F1F5F9')
    pPr.append(shd)

def clean_inline(text):
    """Remove inline markdown formatting."""
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*',     r'\1', text)
    text = re.sub(r'`(.+?)`',       r'\1', text)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text

def convert_md_to_docx(md_path, docx_path, title):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # Cover heading
    cover = doc.add_heading(title, level=0)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cover.runs:
        run.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
        run.font.size = Pt(22)

    sub = doc.add_paragraph('EduTrack LMS — South African Multi-School Learner Management System')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    doc.add_paragraph()

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Code block
        if line.strip().startswith('```'):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # Headings
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text  = clean_inline(m.group(2))
            h = doc.add_heading(text, level=min(level, 4))
            set_heading_style(h, level)
            i += 1
            continue

        # Table
        if line.strip().startswith('|') and i + 1 < len(lines) and lines[i+1].strip().startswith('|'):
            i = add_table_from_md(doc, lines, i)
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}$', line.strip()):
            doc.add_paragraph('─' * 80).runs[0].font.color.rgb = RGBColor(0xd1, 0xd5, 0xdb)
            i += 1
            continue

        # Bullet list
        m_ul = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if m_ul:
            indent = len(m_ul.group(1)) // 2
            text = clean_inline(m_ul.group(2))
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Cm(indent * 0.5)
            p.add_run(text)
            i += 1
            continue

        # Numbered list
        m_ol = re.match(r'^\s*\d+\.\s+(.*)', line)
        if m_ol:
            text = clean_inline(m_ol.group(1))
            p = doc.add_paragraph(style='List Number')
            p.add_run(text)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Normal paragraph
        text = clean_inline(line)
        doc.add_paragraph(text)
        i += 1

    doc.save(docx_path)
    print(f'  Saved: {os.path.basename(docx_path)}')


# ── Convert all docs ───────────────────────────────────────────────────────────

files = [
    ('PROJECT_PLAN.md',    'PROJECT_PLAN.docx',    'Project Plan v2.0'),
    ('DATABASE_DESIGN.md', 'DATABASE_DESIGN.docx', 'Database Design'),
    ('README.md',          'README.docx',           'EduTrack LMS — Overview'),
]

print('Converting documents to Word format...')
for md_name, docx_name, title in files:
    md_path   = os.path.join(BASE, md_name)
    docx_path = os.path.join(BASE, docx_name)
    if os.path.exists(md_path):
        convert_md_to_docx(md_path, docx_path, title)
    else:
        print(f'  Skipped (not found): {md_name}')

print('Done.')
