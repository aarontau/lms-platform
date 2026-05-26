"""
LMS Platform — Module Process Specification Document Generator
Modelled on Academia ERP Process Selection Document (Serosoft, Nov 2025)
Adapted for: Next.js + NestJS + PostgreSQL / South African CAPS context
Author: Dr. B.A. Tau | 26 May 2026
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
from datetime import datetime

# ─── Colour Palette ──────────────────────────────────────────────────────────
NAVY      = RGBColor(0x1E, 0x3A, 0x5F)   # #1E3A5F  — primary brand
INDIGO    = RGBColor(0x4F, 0x46, 0xE5)   # #4F46E5  — primary-600
STEEL     = RGBColor(0x47, 0x5C, 0x78)   # #475C78  — mid tone
SILVER    = RGBColor(0xF1, 0xF5, 0xF9)   # #F1F5F9  — table header fill
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREEN     = RGBColor(0x05, 0x96, 0x69)   # BUILT
AMBER     = RGBColor(0xD9, 0x77, 0x06)   # IN PROGRESS
BLUE      = RGBColor(0x1D, 0x4E, 0xD8)   # PLANNED
GREY      = RGBColor(0x6B, 0x72, 0x80)   # DESCOPED
RED       = RGBColor(0xDC, 0x26, 0x26)   # DESCOPED


# ─── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set cell background colour via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), val.get('val', 'single'))
            el.set(qn('w:sz'), str(val.get('sz', 4)))
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def para_fmt(para, left=0, space_before=0, space_after=6, line_spacing=None):
    pf = para.paragraph_format
    pf.left_indent = Pt(left)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)


def heading(doc, text, level=1, color=NAVY):
    p = doc.add_paragraph()
    para_fmt(p, space_before=10, space_after=4)
    run = p.add_run(text)
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
    elif level == 2:
        run.font.size = Pt(13)
    elif level == 3:
        run.font.size = Pt(11)
    run.font.color.rgb = color
    return p


def body(doc, text, bold=False, italic=False, color=None, size=10, left=0,
         space_before=0, space_after=4):
    p = doc.add_paragraph()
    para_fmt(p, left=left, space_before=space_before, space_after=space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    para_fmt(p, space_before=1, space_after=2)
    p.paragraph_format.left_indent = Pt(18 + level * 12)
    if bold_prefix:
        run = p.add_run(bold_prefix + ' ')
        run.bold = True
        run.font.size = Pt(10)
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def divider(doc, color='1E3A5F'):
    p = doc.add_paragraph()
    para_fmt(p, space_before=2, space_after=2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def page_break(doc):
    doc.add_page_break()


def status_badge(doc, status: str):
    """Insert a coloured status label."""
    colors = {
        'BUILT':        (GREEN,  '1E6B4A'),
        'IN PROGRESS':  (AMBER,  'B45309'),
        'PLANNED':      (BLUE,   '1D4ED8'),
        'DESCOPED':     (GREY,   '6B7280'),
        'PHASE 2':      (BLUE,   '1D4ED8'),
        'PHASE 3':      (STEEL,  '475C78'),
    }
    color, hex_val = colors.get(status, (GREY, '6B7280'))
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, hex_val)
    cell.width = Inches(1.4)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'  {status}  ')
    run.bold = True
    run.font.size = Pt(8)
    run.font.color.rgb = WHITE
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return tbl


def nav_table(doc, rows_data, title=''):
    """2-column table: label | value"""
    if title:
        body(doc, title, bold=True, color=STEEL, size=9, space_before=4, space_after=2)
    tbl = doc.add_table(rows=len(rows_data), cols=2)
    tbl.style = 'Table Grid'
    col_widths = [Inches(1.8), Inches(4.6)]
    for i, (label, value) in enumerate(rows_data):
        row = tbl.rows[i]
        row.height = Pt(16)
        lc = row.cells[0]
        vc = row.cells[1]
        lc.width = col_widths[0]
        vc.width = col_widths[1]
        set_cell_bg(lc, 'EEF2FF')
        lp = lc.paragraphs[0]
        lp.paragraph_format.space_before = Pt(1)
        lp.paragraph_format.space_after = Pt(1)
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = NAVY
        vp = vc.paragraphs[0]
        vp.paragraph_format.space_before = Pt(1)
        vp.paragraph_format.space_after = Pt(1)
        vr = vp.add_run(value)
        vr.font.size = Pt(9)
    return tbl


def module_header(doc, module_name, phase, status):
    """Full-width navy bar with module name."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    row = tbl.rows[0]
    lc = row.cells[0]
    rc = row.cells[1]
    lc.width = Inches(5.2)
    rc.width = Inches(1.2)
    set_cell_bg(lc, '1E3A5F')
    set_cell_bg(rc, '1E3A5F')
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(4)
    lp.paragraph_format.space_after = Pt(4)
    lr = lp.add_run(f'  {module_name}')
    lr.bold = True
    lr.font.size = Pt(13)
    lr.font.color.rgb = WHITE
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp.paragraph_format.space_before = Pt(4)
    rp.paragraph_format.space_after = Pt(4)
    # Phase label
    rr = rp.add_run(phase)
    rr.font.size = Pt(8)
    rr.bold = True
    phase_colors = {
        'Phase 0': 'A7F3D0',
        'Phase 1': 'FDE68A',
        'Phase 2': 'BFDBFE',
        'Phase 3': 'E5E7EB',
    }
    lc.width = Inches(5.0)
    rc.width = Inches(1.4)
    return tbl


def toc_table(doc, modules):
    """Table of contents table."""
    tbl = doc.add_table(rows=1 + len(modules), cols=3)
    tbl.style = 'Table Grid'
    # Header
    hrow = tbl.rows[0]
    for ci, txt in enumerate(['#', 'Module', 'Status']):
        cell = hrow.cells[ci]
        set_cell_bg(cell, '1E3A5F')
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(txt)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = WHITE
    hrow.cells[0].width = Inches(0.4)
    hrow.cells[1].width = Inches(4.8)
    hrow.cells[2].width = Inches(1.2)

    status_colors = {
        'BUILT':       'D1FAE5',
        'IN PROGRESS': 'FEF3C7',
        'PLANNED':     'DBEAFE',
        'PHASE 2':     'DBEAFE',
        'PHASE 3':     'F3F4F6',
        'DESCOPED':    'F3F4F6',
    }
    status_text_colors = {
        'BUILT':       '065F46',
        'IN PROGRESS': '92400E',
        'PLANNED':     '1E40AF',
        'PHASE 2':     '1E40AF',
        'PHASE 3':     '374151',
        'DESCOPED':    '6B7280',
    }

    for i, (num, name, status) in enumerate(modules):
        row = tbl.rows[i + 1]
        row.cells[0].width = Inches(0.4)
        row.cells[1].width = Inches(4.8)
        row.cells[2].width = Inches(1.2)

        bg = 'FFFFFF' if i % 2 == 0 else 'F8FAFC'
        set_cell_bg(row.cells[0], bg)
        set_cell_bg(row.cells[1], bg)
        set_cell_bg(row.cells[2], status_colors.get(status, 'F3F4F6'))

        for ci, txt in enumerate([str(num), name]):
            p = row.cells[ci].paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(txt)
            r.font.size = Pt(9)
            if ci == 1:
                r.font.color.rgb = NAVY

        sp = row.cells[2].paragraphs[0]
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sp.paragraph_format.space_before = Pt(1)
        sp.paragraph_format.space_after = Pt(1)
        sr = sp.add_run(status)
        sr.font.size = Pt(8)
        sr.bold = True
        tc = status_text_colors.get(status, '6B7280')
        sr.font.color.rgb = RGBColor(int(tc[0:2], 16), int(tc[2:4], 16), int(tc[4:6], 16))
    return tbl


# ─── Document Setup ───────────────────────────────────────────────────────────

doc = Document()

# Page margins
from docx.shared import Mm
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(0.8)
section.bottom_margin = Inches(0.8)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — COVER
# ═══════════════════════════════════════════════════════════════════════════════

# Title block
tbl = doc.add_table(rows=1, cols=1)
cell = tbl.cell(0, 0)
set_cell_bg(cell, '1E3A5F')
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
p.paragraph_format.space_after  = Pt(8)
r = p.add_run('LMS PLATFORM')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = WHITE

p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(0)
p2.paragraph_format.space_after  = Pt(6)
r2 = p2.add_run('Module Process Specification Document')
r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(0xA5, 0xB4, 0xFC)  # indigo-300

p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(0)
p3.paragraph_format.space_after  = Pt(24)
r3 = p3.add_run('Version 1.0')
r3.font.size = Pt(10); r3.font.color.rgb = RGBColor(0xC7, 0xD2, 0xFE)  # indigo-200

doc.add_paragraph()

# Document info table
nav_table(doc, [
    ('Document Title',    'LMS Platform — Module Process Specification Document'),
    ('Version',           'V 1.0'),
    ('Date',              '26 May 2026'),
    ('Document Owner',    'Dr. B.A. Tau'),
    ('Prepared For',      'UL Junior, Polokwane (Pilot School)'),
    ('Platform',          'Multi-School SaaS — South African CAPS context'),
    ('Tech Stack',        'Next.js 14 · NestJS · PostgreSQL · Prisma · AWS af-south-1'),
    ('Pilot Launch',      '2 August 2026'),
    ('Full Launch',       'October 2026'),
])

doc.add_paragraph()
body(doc,
     'Prepared by: Dr. B.A. Tau | DMSTE · SoE · FoH · University of Limpopo',
     bold=True, color=NAVY, space_before=8)
body(doc,
     'This document is confidential. It defines the scope, processes, and technical '
     'architecture of the LMS Platform, a South African CAPS-native multi-school '
     'Learning Management System. Its structure is modelled on industry-standard '
     'ERP process selection documentation.', size=9, color=STEEL)

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — INTRODUCTION LETTER
# ═══════════════════════════════════════════════════════════════════════════════

heading(doc, 'Introduction', level=2, color=NAVY)
divider(doc)
doc.add_paragraph()

body(doc, 'To:', bold=True)
body(doc, 'The Management Team')
body(doc, 'UL Junior, Polokwane')
doc.add_paragraph()

body(doc,
     'It is with great pleasure that I present the LMS Platform Module Process '
     'Specification Document for UL Junior and associated pilot schools. This '
     'document captures the full scope of the LMS platform — a purpose-built, '
     'CAPS-native school management system designed for South African public and '
     'independent schools from Grade 4 to Grade 12.')
doc.add_paragraph()
body(doc,
     'The platform is architected as a multi-school SaaS solution, hosted on '
     'AWS Cape Town (af-south-1) to ensure POPI compliance. Every module has '
     'been designed against the South African Schools Act, CAPS curriculum '
     'requirements, and the LURITS/SA-SAMS data standards.')
doc.add_paragraph()
body(doc,
     'This document outlines forty-one modules organised across four delivery '
     'phases. For each module, the Business Case, selected processes, data '
     'models, API surface, and frontend routes are documented. Module status '
     'is indicated as: BUILT · IN PROGRESS · PLANNED · PHASE 2 · PHASE 3 · DESCOPED.')
doc.add_paragraph()
body(doc,
     'The pilot launch is targeted for 2 August 2026, with the full multi-school '
     'commercial launch scheduled for October 2026.', bold=False)
doc.add_paragraph()
body(doc, 'Yours sincerely,', space_after=14)
body(doc, 'Dr. B.A. Tau', bold=True)
body(doc, 'Project Owner & Lead Developer')
body(doc, 'LMS Platform | DMSTE | University of Limpopo')

page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — SCOPE TABLE
# ═══════════════════════════════════════════════════════════════════════════════

heading(doc, 'Scope of Modules', level=2, color=NAVY)
divider(doc)
body(doc,
     'The scope of this document covers the following modules. Descoped items '
     'were excluded after assessment against the South African school management '
     'context and the pilot school profile.', size=9, color=STEEL, space_after=8)

modules = [
    # num, name, status
    (1,  'Organisation / School Setup',              'BUILT'),
    (2,  'User & Roles Management',                  'BUILT'),
    (3,  'Calendar & Academic Year Management',      'BUILT'),
    (4,  'Grade Management',                         'BUILT'),
    (5,  'Learner Information System',               'BUILT'),
    (6,  'Subject Management',                       'IN PROGRESS'),
    (7,  'Timetable Management',                     'PLANNED'),
    (8,  'Attendance Management',                    'PLANNED'),
    (9,  'Homework & Assignment Management',         'PLANNED'),
    (10, 'Assessment Engine (POA + SBA Calculator)', 'PLANNED'),
    (11, 'Teacher Portal',                           'PLANNED'),
    (12, 'Learner Portal',                           'PHASE 2'),
    (13, 'Parent / Guardian Portal',                 'PHASE 2'),
    (14, 'Learner Mobile App',                       'PHASE 3'),
    (15, 'Parent / Guardian Mobile App',             'PHASE 3'),
    (16, 'Teacher Mobile App',                       'PHASE 3'),
    (17, 'Messaging & Notifications',                'PHASE 2'),
    (18, 'Admission Management',                     'PHASE 2'),
    (19, 'Applicant Portal',                         'PHASE 3'),
    (20, 'Report Card & Promotion Management',       'PLANNED'),
    (21, 'LURITS / SA-SAMS Export',                  'PLANNED'),
    (22, 'CAPS Curriculum Engine',                   'IN PROGRESS'),
    (23, 'Programme of Assessment (POA)',             'PLANNED'),
    (24, 'Diagnostic Assessment & Learning Profile', 'PLANNED'),
    (25, 'Fee & Finance Management',                 'PHASE 2'),
    (26, 'Disciplinary Management',                  'PHASE 2'),
    (27, 'Service Request Management',               'PHASE 3'),
    (28, 'Certificate & Document Management',        'PHASE 3'),
    (29, 'Grievance Management',                     'PHASE 3'),
    (30, 'Event Management',                         'PHASE 3'),
    (31, 'Employee / Staff Management',              'PHASE 3'),
    (32, 'Task Management',                          'PHASE 3'),
    (33, 'Analytics & Reporting Dashboard',          'PHASE 2'),
    (34, 'Campaign Management',                      'PHASE 3'),
    (35, 'Front Desk Management',                    'DESCOPED'),
    (36, 'Hostel / Boarding Management',             'DESCOPED'),
    (37, 'Transport Management',                     'DESCOPED'),
    (38, 'Medical Care Management',                  'DESCOPED'),
    (39, 'Security Gate Management',                 'DESCOPED'),
    (40, 'Scholarship Management',                   'DESCOPED'),
    (41, 'Committee Management',                     'DESCOPED'),
]

toc_table(doc, modules)
page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
# MODULE SECTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def module_section(doc, num, name, phase, status,
                   business_case,
                   processes,
                   caps_context=None,
                   models=None,
                   api_prefix=None,
                   frontend_route=None,
                   notes=None):
    """Render a full module section."""
    module_header(doc, f'{num}. {name}', phase, status)
    doc.add_paragraph()

    # Status badge row
    status_badge(doc, status)
    doc.add_paragraph()

    # Business Case
    heading(doc, 'Business Case', level=3, color=STEEL)
    body(doc, business_case, size=10)

    # Process Selected
    heading(doc, 'Process Selected', level=3, color=STEEL)
    for proc_heading, proc_items in processes:
        if proc_heading:
            body(doc, proc_heading, bold=True, size=10, space_after=2)
        for item in proc_items:
            bullet(doc, item)

    # CAPS Context
    if caps_context:
        heading(doc, 'South African / CAPS Context', level=3, color=STEEL)
        for line in caps_context:
            bullet(doc, line)

    # Technical Reference
    tech_rows = []
    if models:
        tech_rows.append(('Prisma Models', ', '.join(models)))
    if api_prefix:
        tech_rows.append(('API Prefix', api_prefix))
    if frontend_route:
        tech_rows.append(('Frontend Route', frontend_route))
    if tech_rows:
        heading(doc, 'Technical Reference', level=3, color=STEEL)
        nav_table(doc, tech_rows)

    # Notes
    if notes:
        doc.add_paragraph()
        body(doc, f'Note: {notes}', italic=True, color=STEEL, size=9)

    divider(doc)
    doc.add_paragraph()


# ── 1. Organisation / School Setup ───────────────────────────────────────────
module_section(doc, 1, 'Organisation / School Setup', 'Phase 0', 'BUILT',
    business_case=(
        'Each school on the platform requires a complete organisational profile before any '
        'academic data can be captured. The school profile is the root tenant record from '
        'which all data is scoped.'
    ),
    processes=[
        ('School Profile', [
            'Create school with full name, EMIS number, type (Public / Independent / IEB / Combined)',
            'Assign province, district, and circuit from the national geo-hierarchy',
            'Upload school logo and configure subdomain (e.g. uljr.lms.co.za)',
            'Set curriculum type: CAPS (default) or IEB',
            'Set school status: TRIAL → ACTIVE → SUSPENDED',
        ]),
        ('Academic Year Configuration', [
            'Create academic year with start and end dates',
            'Mark one academic year as current',
            'Configure terms within academic year (4 terms standard)',
            'Each term has start and end date, term number, and active flag',
        ]),
        ('Grade Configuration', [
            'Create grades (Grade 4 through Grade 12)',
            'Each grade is linked to a CAPS Phase (Intermediate / Senior / FET)',
            'CAPS phase determines SBA:Exam weighting — hard-coded, not configurable',
        ]),
        ('Venue Management', [
            'Create venues: classrooms, laboratories, computer labs, halls, sports fields',
            'Set venue capacity and type',
            'Venues are available for timetable slot assignment',
        ]),
    ],
    caps_context=[
        'Province → District → Circuit hierarchy mirrors the DBE EMIS administrative structure',
        'CAPS phases are pre-loaded and map directly to grade ranges (Gr4–6, Gr7–9, Gr10–11, Gr12)',
        'EMIS number is a required unique identifier for LURITS export compatibility',
    ],
    models=['School', 'AcademicYear', 'Term', 'Grade', 'Venue', 'Province', 'District', 'Circuit', 'CapsPhase'],
    api_prefix='/api/schools',
    frontend_route='/(dashboard)/schools/onboarding',
)

# ── 2. User & Roles ───────────────────────────────────────────────────────────
module_section(doc, 2, 'User & Roles Management', 'Phase 0', 'BUILT',
    business_case=(
        'All staff, parents, and learners who interact with the platform require a user '
        'account with appropriate role-based access. Access is strictly scoped to the '
        'user\'s own school (multi-tenancy). Seven roles cover all user types.'
    ),
    processes=[
        ('User Management', [
            'Create users with first name, last name, email, phone, and role',
            'Users are scoped to a school via school_id (multi-tenant isolation)',
            'Deactivate users without deletion (soft delete)',
            'Track last login date and email verification status',
        ]),
        ('Role Management', [
            'Seven roles: SUPER_ADMIN · SCHOOL_ADMIN · PRINCIPAL · HOD · TEACHER · PARENT · LEARNER',
            'SUPER_ADMIN has cross-school platform access (Serosoft equivalent)',
            'PRINCIPAL and HOD have school-wide visibility of marks and reports',
            'TEACHER has access only to their own classes and subjects',
            'PARENT sees only their linked child\'s data',
            'LEARNER sees only their own timetable, marks, and assignments',
        ]),
        ('Authentication', [
            'JWT-based authentication with school_id embedded in token',
            'NextAuth.js handles session management on the frontend',
            'Password hashed with bcrypt; reset via email link',
            'Token expiry: 1 hour access token, 7-day refresh token',
        ]),
    ],
    caps_context=[
        'HOD role is specific to the South African school context — Head of Department signs off marks',
        'Class teacher is a separate relationship (classTeacherId on Class model) not a role',
    ],
    models=['User'],
    api_prefix='/api/users, /api/auth',
    frontend_route='/(dashboard)/users',
)

# ── 3. Calendar & Academic Year Management ───────────────────────────────────
module_section(doc, 3, 'Calendar & Academic Year Management', 'Phase 0', 'BUILT',
    business_case=(
        'The academic calendar is the temporal backbone of the system. Attendance registers, '
        'assessment tasks, SBA calculations, and report cards all reference the academic '
        'year and term structure.'
    ),
    processes=[
        ('Academic Year', [
            'Create academic years with start date, end date, and year number',
            'One academic year is marked as current at any time',
            'Academic year is a required foreign key on all academic data models',
        ]),
        ('Term Configuration', [
            'Create up to 4 terms per academic year (South African standard)',
            'Each term has a name, start date, end date, and active flag',
            'Attendance registers and assessment tasks reference specific terms',
            'SBA calculations are aggregated per term and annually',
        ]),
    ],
    caps_context=[
        'South African schools operate on a 4-term calendar (Jan–Mar, Apr–Jun, Jul–Sep, Oct–Nov)',
        'CAPS assessment cycles align to terms: one SBA cycle per term per subject',
    ],
    models=['AcademicYear', 'Term'],
    api_prefix='/api/academic-years',
    frontend_route='/(dashboard)/schools/onboarding (term config step)',
)

# ── 4. Grade Management ───────────────────────────────────────────────────────
module_section(doc, 4, 'Grade Management', 'Phase 0', 'BUILT',
    business_case=(
        'Grades define the year-level groups within a school. Each grade is linked to a '
        'CAPS phase, which determines assessment weightings. Classes (sections) are created '
        'within grades. Learners are enrolled into a specific class each academic year.'
    ),
    processes=[
        ('Grade Setup', [
            'Create grades for each academic year offered at the school',
            'Link each grade to its CAPS phase (auto-derived from grade number)',
            'CAPS Phase mapping: Gr1–3 = Foundation, Gr4–6 = Intermediate, Gr7–9 = Senior, Gr10–12 = FET',
        ]),
        ('Class (Section) Management', [
            'Create classes within a grade: e.g. Grade 8A, Grade 8B',
            'Assign a class teacher to each class',
            'Set maximum capacity per class (default: 35)',
            'Classes are the unit of attendance register and homeroom reporting',
        ]),
    ],
    caps_context=[
        'Platform initially covers Grades 4–9 (Intermediate and Senior Phase)',
        'FET Phase (Gr10–12) added in v2 with NSC subject combination validation',
        'Each grade-academic year combination is unique per school',
    ],
    models=['Grade', 'Class', 'CapsPhase'],
    api_prefix='/api/grades',
    frontend_route='/(dashboard)/schools/onboarding (grade config step)',
)

# ── 5. Learner Information System ────────────────────────────────────────────
module_section(doc, 5, 'Learner Information System', 'Phase 0', 'BUILT',
    business_case=(
        'Full biographical, demographic, and academic profile management for each learner. '
        'This is the central learner record from which all academic, attendance, and report '
        'data is derived. Compliance with LURITS data requirements is built in from day one.'
    ),
    processes=[
        ('Learner Registration', [
            'Capture full biographical data: name, date of birth, gender, nationality, home language',
            'Capture identity information: SA ID / passport / birth certificate number',
            'Record admission date, admission number, and previous school',
            'Flag learners with special educational needs (LSEN) and capture medical notes',
            'Auto-generate a unique student number per school on registration',
            'Store LURITS number for SA-SAMS synchronisation',
        ]),
        ('Guardian / Parent Linking', [
            'Create guardian records with name, ID number, phone, email, and relationship',
            'Link multiple guardians to a learner (mother, father, guardian, other)',
            'Designate one guardian as primary contact',
            'Flag guardians who are authorised to collect the learner',
            'Link a guardian to a parent User account for portal access',
        ]),
        ('Learner Enrolment', [
            'Enrol learner into a specific grade and class for an academic year',
            'One active enrolment per learner per academic year',
            'Enrolment statuses: ACTIVE · TRANSFERRED_OUT · WITHDRAWN · COMPLETED',
            'Enrolment drives subject allocation, attendance, and report generation',
        ]),
        ('Learner List & Search', [
            'Filter learners by grade, class, status, gender, home language',
            'Bulk import via CSV (validate before import, preview errors)',
            'Export learner list to Excel or LURITS-compatible format',
            'Print learner ID cards from profile view',
        ]),
    ],
    caps_context=[
        'Home language is a LURITS required field — 11 official SA languages supported',
        'LURITS number is the national learner identifier used in SA-SAMS',
        'Student number auto-format: [SCHOOL_CODE]-[YEAR]-[SEQUENCE], e.g. HRT-2026-001',
        'Special needs (LSEN) flag triggers ISSP/IEP tracking module in Phase 3',
    ],
    models=['Learner', 'LearnerGuardian', 'LearnerEnrolment', 'Guardian'],
    api_prefix='/api/learners',
    frontend_route='/(dashboard)/learners',
)

# ── 6. Subject Management ────────────────────────────────────────────────────
module_section(doc, 6, 'Subject Management', 'Phase 1', 'IN PROGRESS',
    business_case=(
        'Each school selects subjects from the national CAPS catalogue and configures '
        'which subjects are offered per grade. Teachers are assigned to subject-class '
        'combinations. This is the foundation for timetabling and assessment.'
    ),
    processes=[
        ('CAPS Subject Catalogue', [
            'National CAPS subject catalogue is pre-seeded (not editable per school)',
            'Subjects are categorised: Language (HL/FAL/SAL), Mathematics, Sciences, '
            'Social Sciences, Technology, Arts, EMS, Life Orientation',
            'Each subject is tagged as compulsory or optional per CAPS phase',
            'IEB subject catalogue added in Phase 2',
        ]),
        ('School Subject Configuration', [
            'School administrator selects subjects offered at the school from the CAPS catalogue',
            'Custom display name and code can be set per school per subject',
            'Subjects can be activated or deactivated per academic year',
        ]),
        ('Subject-Class Assignment (Subject Class)', [
            'Link a school subject to a specific class for a specific academic year',
            'Assign one teacher to each subject-class combination',
            'Subject classes drive timetable slots, attendance registers, and assessment tasks',
            'One subject can be taught by different teachers in different classes',
        ]),
        ('Learner Subject Enrolment', [
            'Learners are enrolled into subjects via their class enrolment (mandatory subjects auto-enrolled)',
            'Optional subjects selected individually per learner',
            'Subject enrolment is the prerequisite for mark capture',
        ]),
    ],
    caps_context=[
        'CAPS mandates specific compulsory subjects per phase — these are enforced by the platform',
        'Grade 10–12 subject combination validation (minimum NSC requirements) added in v2',
        'Home Language must be one of the 11 official South African languages',
        'Life Orientation is compulsory for all grades in all phases',
    ],
    models=['CapsSubject', 'SchoolSubject', 'SubjectClass'],
    api_prefix='/api/subjects',
    frontend_route='/(dashboard)/subjects',
)

# ── 7. Timetable Management ───────────────────────────────────────────────────
module_section(doc, 7, 'Timetable Management', 'Phase 1', 'PLANNED',
    business_case=(
        'A school timetable maps each subject class to a period, day, teacher, and venue. '
        'The timetable is the schedule backbone for attendance marking (period-by-period) '
        'and teacher workload reporting.'
    ),
    processes=[
        ('Period Configuration', [
            'Define school periods: period number, name (e.g. Period 1), start time, end time, day of week',
            'Mark periods as lessons or non-lessons (break, lunch, assembly)',
            'Periods are configured per academic year',
        ]),
        ('Master Timetable', [
            'Assign each subject class to a period, day, and venue to create a timetable slot',
            'Conflict detection: flag teacher double-booking and venue clashes in real time',
            'View master timetable as a grid: rows = periods, columns = days',
            'Filter timetable by class, teacher, or venue',
        ]),
        ('Timetable Views', [
            'Class timetable view: full week grid per class',
            'Teacher timetable view: all classes and subjects per teacher',
            'Venue timetable view: room occupancy per period',
            'Export timetable to PDF (class and teacher formats)',
        ]),
        ('Timetable Import', [
            'Import master timetable via Excel template for schools with existing timetables',
            'Validation on import: check for conflicts before committing',
        ]),
    ],
    caps_context=[
        'South African schools typically run 7–8 periods per day, 5 days per week',
        'Timetable slot data enables period-by-period attendance in Phase 2',
        'Teacher period loads are used for HR workload reporting in Phase 3',
    ],
    models=['Period', 'TimetableSlot', 'Venue', 'SubjectClass'],
    api_prefix='/api/timetable',
    frontend_route='/(dashboard)/timetable',
    notes='Auto-scheduling (conflict resolver) is a Phase 2 feature. v1 is manual grid entry with conflict flagging only.',
)

# ── 8. Attendance Management ─────────────────────────────────────────────────
module_section(doc, 8, 'Attendance Management', 'Phase 1', 'PLANNED',
    business_case=(
        'Attendance is a statutory requirement under the South African Schools Act. '
        'Absenteeism data must be captured daily and submitted via LURITS annually. '
        'The platform captures daily class-level attendance in v1, with period-by-period '
        'attendance added in Phase 2.'
    ),
    processes=[
        ('Daily Attendance Register', [
            'Class teacher marks attendance for their homeroom class each day',
            'Statuses: PRESENT · ABSENT · LATE · EXCUSED ABSENT',
            'Bulk-mark all present, then flag absences individually',
            'Add notes per learner for absence reason',
            'One register per class per date (unique constraint)',
        ]),
        ('Absence Notifications', [
            'Automatic email notification to primary guardian when learner marked absent',
            'Notification triggered on register capture completion',
            'Notification log stored for audit purposes',
        ]),
        ('Attendance Reports', [
            'Daily attendance summary: class view with present/absent count and percentage',
            'Learner attendance history: full year view, term filter',
            'Learner attendance percentage: flag learners below 80% threshold',
            'Class attendance register: downloadable PDF per class per term',
            'Pending register report: identify teachers who have not captured attendance',
        ]),
        ('LURITS Export', [
            'Attendance data feeds into the annual LURITS export package',
            'Export format: number of school days, present days, absent days per learner',
        ]),
    ],
    caps_context=[
        'SASA Section 29: school must keep attendance records and report to district',
        'LURITS requires annual attendance data per learner in specified format',
        '80% attendance threshold is the statutory minimum — at-risk flag triggers parent notification',
        'Period-by-period attendance (linked to timetable slots) added in Phase 2 per academia ERP model',
    ],
    models=['AttendanceRegister', 'AttendanceRecord', 'Notification'],
    api_prefix='/api/attendance',
    frontend_route='/(dashboard)/attendance',
)

# ── 9. Homework & Assignment Management ──────────────────────────────────────
module_section(doc, 9, 'Homework & Assignment Management', 'Phase 1', 'PLANNED',
    business_case=(
        'Teachers set homework and assignments for learners as part of the School-Based '
        'Assessment programme. Assignments can be tracked via the teacher portal and '
        'submitted via the learner portal in Phase 2.'
    ),
    processes=[
        ('Assignment Creation', [
            'Teacher creates an assignment with title, description, due date, and subject class',
            'Assign to: all learners in class, a group, or individual learner',
            'Assignment types mirror CAPS task types: Homework, Assignment, Oral, Practical',
            'Attach files or documents to the assignment (Phase 2 — S3 file storage)',
        ]),
        ('Assignment Submission (Phase 2)', [
            'Learner submits assignment from learner portal or mobile app',
            'Teacher marks submission as received / not received',
            'If linked to a POA assessment task, marks captured via the Assessment Engine',
        ]),
    ],
    models=['AssessmentTask'],
    api_prefix='/api/assessment/tasks',
    frontend_route='/(dashboard)/assessment/tasks',
    notes='Homework and Assignment management shares the AssessmentTask model with the Assessment Engine. Submission tracking is Phase 2.',
)

# ── 10. Assessment Engine (POA + SBA) ─────────────────────────────────────────
module_section(doc, 10, 'Assessment Engine — POA + CAPS SBA Calculator', 'Phase 1', 'PLANNED',
    business_case=(
        'The Assessment Engine is the most critical module of the platform. It implements '
        'the CAPS Programme of Assessment (POA) framework, captures all four assessment '
        'task types, and calculates School-Based Assessment (SBA) marks with '
        'phase-correct weightings. One calculation error produces wrong report cards. '
        'This module requires 100%% unit test coverage before any other module builds on it.'
    ),
    processes=[
        ('Programme of Assessment (POA)', [
            'Teacher creates a POA per subject class per term',
            'POA defines the total number of assessment tasks required for the term',
            'POA is submitted to HOD for review and approval',
            'Status: DRAFT → SUBMITTED → APPROVED',
        ]),
        ('Assessment Task Configuration', [
            'Create tasks within a POA: title, task type, maximum mark, SBA weight, due date',
            'Task types: DIAGNOSTIC · CLASS_TEST · ASSIGNMENT · HOMEWORK · ORAL · PRACTICAL · SUMMATIVE_EXAM',
            'Mark summative exams separately (isExam flag = true)',
            'Task status lifecycle: DRAFT → ACTIVE → CLOSED → MODERATED',
        ]),
        ('Mark Capture', [
            'Teacher captures raw mark per learner per assessment task',
            'Validation: mark cannot exceed maximum mark',
            'Flag learner as absent (does not enter zero — impacts SBA calculation differently)',
            'Flag learner as exempted from task',
            'Capture notes per mark entry',
            'HOD moderates submitted marks (MODERATED status)',
        ]),
        ('CAPS SBA Calculation Engine', [
            'Term SBA = weighted average of all task marks for the term',
            'Phase-correct SBA:Exam split (hard-coded, not configurable per school):',
            '  — Intermediate Phase Gr4–6:  SBA 75%  /  Exam 25%%',
            '  — Senior Phase Gr7–9:        SBA 60%  /  Exam 40%%',
            '  — FET Phase Gr10–11:         SBA 40%  /  Exam 60%%',
            '  — FET Phase Gr12 (NSC):      SBA 25%  /  Exam 75%%',
            'Annual result = (SBA average × SBA weight) + (Exam mark × Exam weight)',
            'Performance level derived from final percentage (CAPS 1–7 scale)',
            'At-risk flag: learner with SBA below 40%% mid-term',
        ]),
        ('Mark Export', [
            'Export markbook to Excel per subject class',
            'Export individual learner mark summary to PDF',
        ]),
    ],
    caps_context=[
        'CAPS POA is a statutory document — teachers must submit it before term start',
        'CAPS performance levels: 1 (0–29%%) · 2 (30–39%%) · 3 (40–49%%) · 4 (50–59%%) · 5 (60–69%%) · 6 (70–79%%) · 7 (80–100%%)',
        'SBA weightings are non-negotiable per CAPS policy — no school override is permitted',
        'The examination mark for Gr4–9 is the end-of-year November examination',
        'Moderation by HOD is a statutory requirement before marks are published on report cards',
    ],
    models=['ProgrammeOfAssessment', 'AssessmentTask', 'LearnerMark', 'TermSbaResult', 'AnnualSubjectResult'],
    api_prefix='/api/assessment',
    frontend_route='/(dashboard)/assessment',
    notes='This is the critical path module. Full unit test coverage must be verified before Report Card module is built.',
)

# ── 11. Teacher Portal ───────────────────────────────────────────────────────
module_section(doc, 11, 'Teacher Portal', 'Phase 1', 'PLANNED',
    business_case=(
        'Teachers require a single consolidated view of their professional responsibilities: '
        'classes, attendance registers, mark books, timetable, and POA management. The '
        'teacher portal is role-filtered to show only the teacher\'s own data.'
    ),
    processes=[
        ('Teacher Dashboard', [
            'Today\'s timetable: classes teaching today with room and subject',
            'Pending attendance registers: classes where register has not been captured today',
            'Outstanding marks: assessment tasks past due date with marks not yet captured',
            'At-risk learner count across all teacher\'s classes',
        ]),
        ('My Classes', [
            'List all subject classes assigned to the teacher for the current academic year',
            'View class list per subject (all enrolled learners)',
            'Navigate directly to mark book, attendance, or POA for each class',
        ]),
        ('Mark Book', [
            'Grid view: rows = learners, columns = assessment tasks',
            'Inline mark entry with instant SBA percentage update',
            'Visual flags: absent (grey), exempted (blue), at-risk (red)',
            'Submit mark book to HOD for moderation',
        ]),
        ('My Profile', [
            'View personal details, qualifications, subject assignments',
            'Download own timetable as PDF',
        ]),
    ],
    models=['User', 'SubjectClass', 'Class', 'AssessmentTask', 'LearnerMark', 'AttendanceRegister'],
    api_prefix='/api/teacher (filtered views)',
    frontend_route='/(dashboard)/teacher',
)

# ── 12. Learner Portal ───────────────────────────────────────────────────────
module_section(doc, 12, 'Learner Portal', 'Phase 2', 'PHASE 2',
    business_case=(
        'Learners access a read-only view of their own academic record: timetable, '
        'attendance, marks, assignments, and report cards. This is a Phase 2 feature '
        'delivered at Week 8 of the project plan.'
    ),
    processes=[
        ('Learner Dashboard', [
            'Today\'s timetable from the master timetable',
            'Recent marks: latest 5 assessment results',
            'Attendance percentage for current term',
            'Upcoming assignments due dates',
        ]),
        ('My Academics', [
            'Timetable: full week grid view',
            'Marks: per subject, per term — all tasks and SBA percentage',
            'Attendance: term-by-term attendance percentage',
            'Report cards: view and download published PDF report cards',
        ]),
        ('Requests', [
            'Raise service requests (Phase 3)',
        ]),
    ],
    models=['Learner', 'LearnerEnrolment', 'LearnerMark', 'AttendanceRecord', 'ReportCard'],
    api_prefix='/api/portal/learner',
    frontend_route='/(portal)/learner',
)

# ── 13. Parent / Guardian Portal ─────────────────────────────────────────────
module_section(doc, 13, 'Parent / Guardian Portal', 'Phase 2', 'PHASE 2',
    business_case=(
        'Parents and guardians access a view of their linked child\'s academic record. '
        'This is the primary channel for parent engagement and replaces the need for '
        'physical report card collection in many cases.'
    ),
    processes=[
        ('Parent Dashboard', [
            'Child\'s attendance status for today',
            'Latest marks and assessment results',
            'At-risk subject flags (SBA below 40%%)',
            'Upcoming report card publication notifications',
        ]),
        ('Child\'s Academics', [
            'Timetable for the current week',
            'Full mark history per subject per term',
            'Attendance record: calendar view of present / absent days',
            'Disciplinary records (when recorded by school)',
        ]),
        ('Reports & Documents', [
            'View and download published term and annual report cards as PDF',
            'Receive push / email notification when report card is published',
        ]),
    ],
    models=['Guardian', 'LearnerGuardian', 'Learner', 'ReportCard', 'AttendanceRecord', 'LearnerMark'],
    api_prefix='/api/portal/parent',
    frontend_route='/(portal)/parent',
)

# ── 14–16. Mobile Apps ───────────────────────────────────────────────────────
module_section(doc, 14, 'Learner, Parent & Teacher Mobile Apps', 'Phase 3', 'PHASE 3',
    business_case=(
        'Mobile applications for learners, parents, and teachers provide on-the-go access '
        'to the most frequently used portal functions. Load-shedding resilience (offline '
        'mode with background sync) is a key requirement for South African schools.'
    ),
    processes=[
        ('Shared Features', [
            'React Native + Expo — one codebase for iOS and Android',
            'Offline mode: IndexedDB local cache with background sync when online',
            'Push notifications via Firebase Cloud Messaging',
            'Biometric login (fingerprint / face ID)',
        ]),
        ('Teacher App', [
            'Capture daily attendance register offline',
            'Enter marks for assessment tasks',
            'View class timetable and today\'s schedule',
        ]),
        ('Parent App', [
            'View child\'s attendance and marks',
            'Receive absence alerts and report card notifications',
            'Download report cards as PDF',
        ]),
        ('Learner App', [
            'View timetable and upcoming assessments',
            'View own marks and attendance',
        ]),
    ],
    models=['User', 'AttendanceRegister', 'LearnerMark', 'Notification'],
    api_prefix='/api (same endpoints as web portals)',
    frontend_route='React Native app (separate repo)',
    notes='Phase 3 — post-pilot. Load-shedding offline resilience is non-negotiable for SA market.',
)

# ── 17. Messaging & Notifications ────────────────────────────────────────────
module_section(doc, 17, 'Messaging & Notifications', 'Phase 2', 'PHASE 2',
    business_case=(
        'Automated and ad-hoc notifications to parents, learners, and teachers are a '
        'critical engagement channel. South African schools rely heavily on WhatsApp and '
        'SMS for parent communication. The platform delivers via email (v1), SMS (v2), '
        'and push notifications (Phase 3).'
    ),
    processes=[
        ('Automated Notifications', [
            'Absence alert: email to primary guardian when learner marked absent',
            'Report card ready: email to all parents when report cards published',
            'Mark published: email to parent when SBA marks are released',
            'At-risk alert: email to HOD when learner SBA drops below 40%%',
        ]),
        ('Ad-hoc Messaging', [
            'School administrator sends message to: all parents, all teachers, specific grade, specific class',
            'Message types: EMAIL · IN_APP · SMS (Phase 2) · PUSH (Phase 3)',
            'Message log stored per recipient for audit',
        ]),
        ('Notification Channels', [
            'Email: AWS SES (af-south-1, POPI compliant)',
            'SMS: BulkSMS or Clickatell integration (Phase 2)',
            'Push notifications: Firebase Cloud Messaging (Phase 3)',
        ]),
    ],
    caps_context=[
        'Parent must be notified within 24 hours of learner absence (SASA obligation)',
        'SMS is preferred channel for rural schools — BulkSMS has strong SA coverage',
    ],
    models=['Notification'],
    api_prefix='/api/notifications',
    frontend_route='/(dashboard)/notifications',
)

# ── 18. Admission Management ─────────────────────────────────────────────────
module_section(doc, 18, 'Admission Management', 'Phase 2', 'PHASE 2',
    business_case=(
        'Schools manage new learner applications for each academic year. Applications are '
        'created via the applicant portal or by admin staff. Approved applications convert '
        'into learner registrations.'
    ),
    processes=[
        ('Admission Setup', [
            'Configure admission cycle per grade per academic year',
            'Set application open and close dates',
            'Set seat capacity per grade',
        ]),
        ('Application Process', [
            'Prospective learner / parent submits application via applicant portal',
            'Admin creates application on behalf of applicant',
            'Application review: SUBMITTED → APPROVED / REJECTED',
            'On approval: learner record created automatically from application data',
        ]),
        ('Applicant Portal', [
            'Public-facing registration form — no login required to start',
            'Parent receives login credentials to complete and track application',
            'Upload supporting documents: birth certificate, previous school report',
            'Track application status online',
        ]),
    ],
    models=['Learner', 'LearnerEnrolment'],
    api_prefix='/api/admissions',
    frontend_route='/(dashboard)/admissions',
    notes='Phase 2. Basic admission via admin is available in Phase 1 (direct learner registration). Full applicant portal in Phase 2.',
)

# ── 20. Report Card & Promotion Management ───────────────────────────────────
module_section(doc, 20, 'Report Card & Promotion Management', 'Phase 1', 'PLANNED',
    business_case=(
        'Term and annual report cards are a statutory deliverable for every South African '
        'school. The platform generates CAPS-formatted report cards as PDFs, with '
        'promotion decisions calculated automatically and presented for principal '
        'review and override.'
    ),
    processes=[
        ('Report Card Generation', [
            'Generate term report cards: per learner SBA marks per subject',
            'Generate annual report cards: SBA average + exam mark + final mark + performance level',
            'Performance level displayed per subject (CAPS 1–7 scale)',
            'Teacher and principal remarks fields',
            'Bulk generation: generate all report cards for a class or grade in one action',
        ]),
        ('Report Card Publication', [
            'Principal reviews and publishes report cards (DRAFT → PUBLISHED)',
            'Published report cards are locked — mark changes require re-generation',
            'PDF stored in S3; accessible via parent portal and email link',
            'Email notification sent to all parents on publication',
        ]),
        ('Promotion Engine', [
            'Automatic promotion recommendation calculated per CAPS rules:',
            '  — Intermediate (Gr4–6): 50%% HL, 40%% FAL, 40%% Maths, 40%% NST, 40%% SS',
            '  — Senior (Gr7–9): 50%% HL, 40%% FAL, 40%% Maths, 40%% in 3 others, 30%% rest',
            'Recommendation: PROMOTE · REPEAT · PROGRESS',
            'Principal can override recommendation with mandatory reason',
            'Override is logged to audit trail',
            'Max 2 years in same grade per phase enforced by progression rules',
        ]),
        ('At-Risk Monitoring', [
            'At-risk flag visible on learner profile from mid-term (4 weeks before term end)',
            'HOD at-risk dashboard: count per subject class',
            'Principal at-risk dashboard: school-wide view per grade',
        ]),
    ],
    caps_context=[
        'CAPS promotion criteria are legislated — no school-level override of the pass criteria',
        'CAPS performance levels 1–7 must appear on all report cards',
        'Principal override of promotion decision is permitted but must be documented',
        'Report cards must be issued within 3 weeks of end of term (SASA obligation)',
    ],
    models=['ReportCard', 'PromotionDecision', 'AnnualSubjectResult', 'TermSbaResult'],
    api_prefix='/api/reports',
    frontend_route='/(dashboard)/reports',
)

# ── 21. LURITS / SA-SAMS Export ──────────────────────────────────────────────
module_section(doc, 21, 'LURITS / SA-SAMS Export', 'Phase 1', 'PLANNED',
    business_case=(
        'Every South African school must submit learner data to the DBE annually via the '
        'Learner Unit Record Information and Tracking System (LURITS). The SA-SAMS '
        'software is the national standard. The platform generates export files compatible '
        'with SA-SAMS import, eliminating the need for double data capture.'
    ),
    processes=[
        ('Export Types', [
            'Learner biographical data export (LURITS format)',
            'Annual attendance data export (school days, present days, absent days per learner)',
            'Marks export (end-of-year final marks per subject per learner)',
            'Full EMIS annual survey package',
        ]),
        ('Export Workflow', [
            'Admin selects export type and academic year',
            'Platform validates data before export: flag missing LURITS numbers, incomplete records',
            'Admin reviews validation report and resolves flagged issues',
            'Generate export file (XML / CSV per SA-SAMS spec)',
            'Download export file for upload to SA-SAMS',
            'Export batch logged with record count, status, and timestamp',
        ]),
        ('Export History', [
            'All export batches logged with status: PENDING · PROCESSING · COMPLETE · FAILED',
            'Re-run failed exports without recreating from scratch',
            'Audit trail: who exported what, when',
        ]),
    ],
    caps_context=[
        'LURITS submission is a statutory obligation under the Education Management Information System (EMIS)',
        'SA-SAMS is the DBE\'s national school administration software — export must be compatible',
        'LURITS number is the national learner identifier — schools without it cannot submit',
        'Export must include all learners, including those who transferred out during the year',
    ],
    models=['LuritsExportBatch', 'Learner', 'AttendanceRecord', 'AnnualSubjectResult'],
    api_prefix='/api/lurits',
    frontend_route='/(dashboard)/lurits',
)

# ── 22. CAPS Curriculum Engine ────────────────────────────────────────────────
module_section(doc, 22, 'CAPS Curriculum Engine', 'Phase 0', 'BUILT',
    business_case=(
        'The CAPS Curriculum Engine is a platform-native module with no equivalent in '
        'generic ERP systems. It provides the full national CAPS subject catalogue, '
        'phase-to-grade mappings, and SBA weighting rules as immutable platform data. '
        'This engine is the source of truth for all curriculum decisions.'
    ),
    processes=[
        ('Global Subject Catalogue', [
            'CapsSubject model: all CAPS subjects pre-seeded at database level',
            'Subjects tagged by: phase, group (Language / Maths / Sciences / etc.), compulsory flag',
            'Curriculum type: CAPS or IEB (IEB catalogue added in Phase 2)',
            'School-level overrides are NOT permitted — catalogue is platform-managed',
        ]),
        ('Phase Mappings & Weightings', [
            'CapsPhase model: Foundation (Gr1–3), Intermediate (Gr4–6), Senior (Gr7–9), FET (Gr10–12)',
            'SBA weight and exam weight stored per phase — not per school',
            'Phase is auto-derived from grade number on grade creation',
        ]),
    ],
    caps_context=[
        'CAPS = Curriculum and Assessment Policy Statement (South Africa, 2011)',
        'Replacing the Academia ERP "Outcome Based Education" module (descoped for UL Junior)',
        'NSC subject combination validation (FET Phase) added in Phase 2',
    ],
    models=['CapsPhase', 'CapsSubject'],
    api_prefix='/api/curriculum (read-only)',
    frontend_route='Referenced in subject configuration UI',
)

# ── 25. Fee & Finance Management ─────────────────────────────────────────────
module_section(doc, 25, 'Fee & Finance Management', 'Phase 2', 'PHASE 2',
    business_case=(
        'School fees management including fee structure configuration, invoice generation '
        'per learner per term, and payment recording. Phase 1 payment is manual recording. '
        'Phase 2 adds PayFast (SA payment gateway) integration.'
    ),
    processes=[
        ('Fee Structure', [
            'Configure fee types: TUITION · REGISTRATION · ACTIVITY · SPORT · OTHER',
            'Set billing frequency: ANNUAL · TERMLY · MONTHLY',
            'Set amounts per grade and fee type per academic year',
        ]),
        ('Invoice Generation', [
            'Generate invoices per learner based on fee structure',
            'Invoice statuses: DRAFT · ISSUED · PARTIALLY_PAID · PAID · OVERDUE · CANCELLED',
            'Bulk generate invoices for a full grade at once',
        ]),
        ('Payment Recording', [
            'Manual payment capture (v1): cash, EFT, card, debit order, bursary',
            'Record reference number and generate receipt number',
            'Outstanding debtors report: sortable, export to Excel',
        ]),
        ('PayFast Integration (Phase 2)', [
            'Online payment via PayFast (South African payment gateway)',
            'Parent pays from parent portal — invoice updated automatically',
        ]),
    ],
    models=['FeeStructure', 'Invoice', 'Payment'],
    api_prefix='/api/finance',
    frontend_route='/(dashboard)/finance',
    notes='Fee management is Phase 2. Academic modules take priority for pilot launch.',
)

# ── 26. Disciplinary Management ──────────────────────────────────────────────
module_section(doc, 26, 'Disciplinary Management', 'Phase 2', 'PHASE 2',
    business_case=(
        'Disciplinary incidents are recorded per learner, visible to school staff and '
        'parents on the parent portal. The system supports positive and corrective '
        'behaviour recording in alignment with the school\'s Code of Conduct.'
    ),
    processes=[
        ('Incident Recording', [
            'Record disciplinary incident with: date, type, description, action taken',
            'Link to one or more learners',
            'Positive behaviour recording (merits / commendations) supported alongside incidents',
            'Attach supporting documents (Phase 2 — S3 file storage)',
        ]),
        ('Parent Visibility', [
            'Disciplinary records visible to linked parent on parent portal',
            'School can configure whether parents see all incidents or only serious ones',
        ]),
    ],
    models=['Learner (disciplinary activity sub-record — to be added in Phase 2)'],
    api_prefix='/api/discipline',
    frontend_route='/(dashboard)/discipline',
)

# ── 33. Analytics Dashboard ───────────────────────────────────────────────────
module_section(doc, 33, 'Analytics & Reporting Dashboard', 'Phase 2', 'PHASE 2',
    business_case=(
        'Data-driven decision making at principal, HOD, and district level. The analytics '
        'module provides aggregated views of academic performance, attendance, and learner '
        'risk indicators. No additional data capture — all metrics derive from existing records.'
    ),
    processes=[
        ('Principal Dashboard', [
            'School-wide pass rate per grade and subject',
            'At-risk learner count per grade (SBA below threshold)',
            'Attendance rate: 7-day rolling per grade',
            'Subject performance ranking: best to worst per grade',
            'Report card publication status (how many pending vs published)',
        ]),
        ('HOD Dashboard', [
            'Subject performance per class: average mark, pass rate, mark distribution chart',
            'At-risk count per class: flag learners needing intervention',
            'Teacher mark submission status: pending vs completed per POA',
        ]),
        ('Analytical Reports', [
            'Year-on-year pass rate comparison per subject (Phase 3)',
            'Cohort tracking: follow a group of learners across grades (Phase 3)',
            'LURITS statistical data: gender, home language, nationality breakdowns',
        ]),
    ],
    models=['TermSbaResult', 'AnnualSubjectResult', 'AttendanceRecord', 'ReportCard', 'PromotionDecision'],
    api_prefix='/api/analytics',
    frontend_route='/(dashboard)/analytics',
)

# Descoped modules
def descoped_module(doc, num, name, reason):
    tbl = doc.add_table(rows=1, cols=2)
    lc = tbl.cell(0, 0); rc = tbl.cell(0, 1)
    lc.width = Inches(5.0); rc.width = Inches(1.4)
    set_cell_bg(lc, '6B7280'); set_cell_bg(rc, '6B7280')
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(3)
    lp.paragraph_format.space_after  = Pt(3)
    lr = lp.add_run(f'  {num}. {name}')
    lr.bold = True; lr.font.size = Pt(11); lr.font.color.rgb = WHITE
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp.paragraph_format.space_before = Pt(3)
    rp.paragraph_format.space_after  = Pt(3)
    rr = rp.add_run('DESCOPED')
    rr.bold = True; rr.font.size = Pt(8); rr.font.color.rgb = WHITE
    body(doc, f'Reason: {reason}', italic=True, color=GREY, size=9, space_after=8)
    divider(doc, color='9CA3AF')
    doc.add_paragraph()

heading(doc, 'Descoped Modules', level=2, color=GREY)
body(doc,
     'The following modules were assessed against the South African school management '
     'context and the pilot school profile. They were excluded from all project phases.',
     size=9, color=STEEL, space_after=8)

descoped_module(doc, 35, 'Front Desk Management',
    'Not applicable to typical South African school structure. Managed manually.')
descoped_module(doc, 36, 'Hostel / Boarding Management',
    'UL Junior and pilot schools are day schools. No boarding facilities.')
descoped_module(doc, 37, 'Transport Management',
    'School transport is parent-managed. Not in scope for the LMS platform.')
descoped_module(doc, 38, 'Medical Care Management',
    'Medical records captured in learner profile (special needs / medical notes). '
    'Full medical care module is out of scope for the school management platform.')
descoped_module(doc, 39, 'Security Gate Management',
    'Physical access control is not managed by the LMS. Out of scope.')
descoped_module(doc, 40, 'Scholarship Management',
    'No bursary or scholarship management in pilot schools. Deferred indefinitely.')
descoped_module(doc, 41, 'Committee Management',
    'School governing body (SGB) committees are managed outside the LMS.')


# ═══════════════════════════════════════════════════════════════════════════════
# DEVELOPMENT ROADMAP
# ═══════════════════════════════════════════════════════════════════════════════

page_break(doc)
heading(doc, 'Development Roadmap', level=2, color=NAVY)
divider(doc)
body(doc,
     'The table below maps every module to its development phase, delivery week, '
     'and current build status as of 26 May 2026.', size=9, color=STEEL, space_after=8)

roadmap = [
    # Phase, Module, Week, Status, Built Components, Next Action
    ('0', 'School Setup & Config',          'Wk 1–2', 'BUILT',       'API + UI complete',                   'Monitor feedback from pilot'),
    ('0', 'User & Roles',                   'Wk 2',   'BUILT',       'Auth + RBAC + UI complete',            'Monitor feedback from pilot'),
    ('0', 'Calendar / Academic Year',       'Wk 1–2', 'BUILT',       'API + seed data complete',             'Add term edit UI'),
    ('0', 'Grade Management',               'Wk 2',   'BUILT',       'Grade + Class API + seed complete',    'Grade config UI for admin'),
    ('0', 'Learner Information System',     'Wk 3',   'BUILT',       'Full CRUD + Guardian + Enrolment',     'CSV bulk import UI'),
    ('0', 'CAPS Curriculum Engine',         'Wk 1',   'BUILT',       'Schema seeded with all CAPS subjects', 'Seed IEB catalogue in Phase 2'),
    ('1', 'Subject Management',             'Wk 3',   'IN PROGRESS', 'Schema + API scaffolded',              'Build Subject Management UI'),
    ('1', 'Timetable Management',           'Wk 4',   'PLANNED',     'Schema complete',                      'Build Period + TimetableSlot APIs'),
    ('1', 'Attendance Management',          'Wk 4',   'PLANNED',     'Schema complete',                      'Build Attendance Register API + UI'),
    ('1', 'Homework & Assignments',         'Wk 4–5', 'PLANNED',     'AssessmentTask model ready',           'Build task creation + list UI'),
    ('1', 'Assessment Engine (POA + SBA)',  'Wk 5–6', 'PLANNED',     'Schema complete',                      'Build POA API + SBA calculator + unit tests'),
    ('1', 'Teacher Portal',                 'Wk 8',   'PLANNED',     'Auth role guard ready',                'Build teacher dashboard + mark book'),
    ('1', 'Report Cards + Promotion',       'Wk 7',   'PLANNED',     'ReportCard + PromotionDecision schema', 'Build generation engine + PDF export'),
    ('1', 'LURITS / SA-SAMS Export',        'Wk 8',   'PLANNED',     'LuritsExportBatch schema ready',       'Build export engine + validation'),
    ('2', 'Learner Portal',                 'Wk 8',   'PHASE 2',     'Route structure exists',               'Build portal pages post-pilot'),
    ('2', 'Parent Portal',                  'Wk 8',   'PHASE 2',     'Route structure exists',               'Build portal pages post-pilot'),
    ('2', 'Messaging & Notifications',      'Wk 8',   'PHASE 2',     'Notification schema ready',            'Integrate AWS SES, then SMS'),
    ('2', 'Admission Management',           'Sep 26', 'PHASE 2',     'Learner schema supports admission',    'Build admission cycle + applicant portal'),
    ('2', 'Fee & Finance Management',       'Wk 9',   'PHASE 2',     'FeeStructure + Invoice + Payment',     'Build fee config + invoice generation UI'),
    ('2', 'Disciplinary Management',        'Sep 26', 'PHASE 2',     'Not started',                          'Add DisciplinaryRecord model + UI'),
    ('2', 'Analytics Dashboard',            'Wk 9',   'PHASE 2',     'Not started',                          'Build principal + HOD analytics views'),
    ('3', 'Mobile Apps (3 apps)',           'Q4 26',  'PHASE 3',     'Not started',                          'React Native + Expo scaffold post-launch'),
    ('3', 'SMS Gateway',                    'Aug 26', 'PHASE 3',     'Not started',                          'BulkSMS / Clickatell integration'),
    ('3', 'Staff / HR Management',          'Q4 26',  'PHASE 3',     'User + Employee model',                'Full HR module post-launch'),
    ('3', 'Event Management',               'Q4 26',  'PHASE 3',     'Not started',                          'Post-launch'),
    ('3', 'Service Requests',               'Q4 26',  'PHASE 3',     'Not started',                          'Post-launch'),
    ('3', 'Grievance Management',           'Q4 26',  'PHASE 3',     'Not started',                          'Post-launch'),
    ('3', 'Certificate & Documents',        'Q4 26',  'PHASE 3',     'Not started',                          'Post-launch'),
    ('3', 'Campaign Management',            'Q4 26',  'PHASE 3',     'Not started',                          'Post-launch'),
    ('3', 'PAD Platform Integration',       'Q1 27',  'PHASE 3',     'Not started',                          'Roadmap item — next platform'),
]

# Build roadmap table
rt = doc.add_table(rows=1 + len(roadmap), cols=6)
rt.style = 'Table Grid'
headers = ['Ph', 'Module', 'Week', 'Status', 'Built', 'Next Action']
widths = [Inches(0.35), Inches(2.0), Inches(0.5), Inches(0.9), Inches(1.4), Inches(1.5)]
hrow = rt.rows[0]
for ci, (txt, w) in enumerate(zip(headers, widths)):
    cell = hrow.cells[ci]
    cell.width = w
    set_cell_bg(cell, '1E3A5F')
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(txt)
    r.bold = True; r.font.size = Pt(8); r.font.color.rgb = WHITE

status_bg = {
    'BUILT':       'D1FAE5', 'IN PROGRESS': 'FEF3C7',
    'PLANNED':     'DBEAFE', 'PHASE 2':     'EDE9FE',
    'PHASE 3':     'F3F4F6', 'DESCOPED':    'F3F4F6',
}
phase_bg = {'0': 'EEF2FF', '1': 'FEF9C3', '2': 'EDE9FE', '3': 'F0FDF4'}

for i, (phase, module, week, status, built, nxt) in enumerate(roadmap):
    row = rt.rows[i + 1]
    for ci, w in enumerate(widths):
        row.cells[ci].width = w

    bg = 'FFFFFF' if i % 2 == 0 else 'F8FAFC'
    for ci in range(6):
        set_cell_bg(row.cells[ci], bg)

    set_cell_bg(row.cells[0], phase_bg.get(phase, 'F8FAFC'))
    set_cell_bg(row.cells[3], status_bg.get(status, 'F8FAFC'))

    for ci, txt in enumerate([phase, module, week, status, built, nxt]):
        p = row.cells[ci].paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(txt)
        r.font.size = Pt(8)
        if ci == 1:
            r.font.color.rgb = NAVY
        if ci == 0:
            r.bold = True; r.font.color.rgb = NAVY
        if ci == 3:
            r.bold = True


# ═══════════════════════════════════════════════════════════════════════════════
# SIGN-OFF
# ═══════════════════════════════════════════════════════════════════════════════

page_break(doc)
heading(doc, 'Sign-Off', level=2, color=NAVY)
divider(doc)
doc.add_paragraph()
nav_table(doc, [
    ('Document Prepared',   '26 May 2026'),
    ('Phase 0 Completed',   '25 May 2026'),
    ('Pilot Launch Target', '2 August 2026'),
    ('Full Launch Target',  'October 2026'),
    ('Platform',            'LMS — South African CAPS-Native Multi-School SaaS'),
])
doc.add_paragraph()
body(doc,
     'I, Dr. B.A. Tau, acknowledge that this Module Process Specification Document '
     'accurately reflects the scope, processes, and architecture of the LMS Platform '
     'as of 26 May 2026. All modules listed as BUILT have been implemented, tested, '
     'and seeded with pilot school data.', size=10)
doc.add_paragraph()
body(doc, 'Authorised Signatory', bold=True)
body(doc, 'Dr. B.A. Tau', bold=True, color=NAVY)
body(doc, 'Project Owner | DMSTE · SoE · FoH | University of Limpopo')
doc.add_paragraph()

# Footer note
body(doc,
     'This document is modelled on the Academia ERP Process Selection Document '
     '(Serosoft Solutions, November 2025), adapted for the LMS Platform architecture '
     'and South African legislative context.',
     italic=True, color=GREY, size=8)

# ─── Save ─────────────────────────────────────────────────────────────────────
output_path = r"C:\Users\aaron.tau\Desktop\LMS Project\LMS_Platform_Module_Process_Spec_V1_20260526.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
